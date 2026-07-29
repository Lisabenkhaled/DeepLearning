# -*- coding: utf-8 -*-
"""
Build the Word report for the CAC 40 LSTM/GRU project
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Sequence

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "outputs"
DOCX_PATH = ROOT / "rapport_projet_cac40_lstm_gru.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
TABLE_BORDER = "B8C2CC"


def set_run_font(
    run,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    """
    Apply the report font settings to a run

    :param run: Word run to format
    :param size: optional font size in points
    :param color: optional RGB color
    :param bold: optional bold setting
    :param italic: optional italic setting
    """
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    """
    Apply background shading to a table cell

    :param cell: table cell
    :param fill: hexadecimal fill color without '#'
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 80, bottom: int = 80) -> None:
    """
    Apply explicit cell margins in DXA 

    :param cell: table cell
    :param top: top margin in DXA (80 DXA = 1.4 mm)
    :param bottom: bottom margin in DXA
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "bottom": bottom,
        "start": 120,
        "end": 120,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    """
    Apply restrained single-grid borders to a table

    :param table: Word table
    """
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), TABLE_BORDER)


def set_table_width(table, widths: Sequence[float]) -> None:
    """
    Apply fixed widths to all cells in a table

    :param table: Word table
    :param widths: column widths in inches
    """
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            set_cell_margins(row.cells[index])
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def configure_document(doc: Document) -> None:
    """
    Configure page geometry and report styles

    :param doc: Word document
    """
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "Projet CAC 40 - LSTM / GRU"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Rapport de projet"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    """
    Add the first-page title block

    :param doc: Word document
    """
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Prévision de la volatilité intraday du CAC 40")
    set_run_font(run, size=22, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Analyse statistique, bêta, LSTM et GRU")
    set_run_font(run, size=13, color=MUTED)

    metadata = doc.add_table(rows=4, cols=2)
    set_table_width(metadata, [1.5, 5.0])
    set_table_borders(metadata)
    rows = [
        ("Cours", "Apprentissage machine et mégadonnées en finance"),
        ("Données", "cac40.csv"),
        ("Méthode", "Programmation orientée objet avec pandas, NumPy, TensorFlow/Keras"),
        ("Livrables", "Fichier Python, tableaux de résultats, graphiques et rapport Word"),
    ]
    for row, (label, value) in zip(metadata.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].text = label
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            set_run_font(paragraph.runs[0], bold=True)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_paragraph(doc: Document, text: str) -> None:
    """
    Add a standard body paragraph

    :param doc: Word document
    :param text: paragraph text
    """
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    """
    Add bullet paragraphs using Word's list style

    :param doc: Word document
    :param items: bullet item texts
    """
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_metrics_table(doc: Document, metrics: pd.DataFrame) -> None:
    """
    Add the LSTM/GRU metrics table

    :param doc: Word document
    :param metrics: model metrics
    """
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [1.3, 1.4, 1.4, 2.4])
    set_table_borders(table)
    headers = ["Modèle", "RMSE", "R²", "Temps d'entraînement"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        set_run_font(cell.paragraphs[0].runs[0], bold=True)

    for _, row_data in metrics.iterrows():
        row = table.add_row()
        values = [
            str(row_data["model"]),
            f"{row_data['rmse']:.4f}",
            f"{row_data['r2']:.4f}",
            f"{row_data['training_time_seconds']:.2f} secondes",
        ]
        for cell, value in zip(row.cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def describe_bias(mean_error: float) -> str:
    """
    Describe whether a model tends to overestimate or underestimate

    :param mean_error: average prediction error
    :return: textual interpretation
    """
    if mean_error < 0:
        return "sous-estimer"
    if mean_error > 0:
        return "surestimer"
    return "ne pas présenter de biais moyen marqué"


def add_explicit_answers(
    doc: Document,
    metrics: pd.DataFrame,
    predictions: pd.DataFrame,
) -> None:
    """
    Add a section that answers the assignment questions explicitly.

    :param doc: Word document
    :param metrics: model metrics table
    :param predictions: test predictions and errors
    """
    best = metrics.sort_values("rmse").iloc[0]
    fastest = metrics.sort_values("training_time_seconds").iloc[0]
    lstm_bias = predictions["lstm_error"].mean()
    gru_bias = predictions["gru_error"].mean()
    worst_lstm = predictions.assign(
        absolute_error=predictions["lstm_error"].abs()
    ).nlargest(1, "absolute_error").iloc[0]
    worst_gru = predictions.assign(
        absolute_error=predictions["gru_error"].abs()
    ).nlargest(1, "absolute_error").iloc[0]

    doc.add_heading("7. Réponses explicites aux questions du sujet", level=1)
    add_paragraph(
        doc,
        "Cette section reprend directement les questions posées dans l'énoncé "
        "afin de rendre les réponses facilement identifiables."
    )
    answers = [
        (
            "Quel modèle donne les meilleures performances ? "
            f"Le modèle {best['model']} obtient le RMSE le plus faible "
            f"({best['rmse']:.4f}) et le R² le plus élevé "
            f"({best['r2']:.4f}) lors de l'exécution finale."
        ),
        (
            "Lequel est le plus rapide à entraîner ? "
            f"Le modèle {fastest['model']} est le plus rapide, avec "
            f"{fastest['training_time_seconds']:.2f} secondes d'entraînement."
        ),
        (
            "Observe-t-on un bon alignement entre prédictions et réalité ? "
            "L'alignement est correct pour le niveau moyen de volatilité, mais "
            "les modèles lissent les mouvements extrêmes et réagissent moins bien "
            "aux pics soudains."
        ),
        (
            "Les modèles sous-estiment-ils ou surestiment-ils certaines valeurs ? "
            f"En moyenne, le LSTM tend à {describe_bias(lstm_bias)} "
            f"(erreur moyenne {lstm_bias:.4f}) et le GRU tend à "
            f"{describe_bias(gru_bias)} (erreur moyenne {gru_bias:.4f}). "
            "Les plus grands écarts sont surtout des sous-estimations lors des "
            "journées de forte volatilité."
        ),
        (
            "Quelles sont les périodes où les erreurs sont les plus fortes ? "
            f"Pour le LSTM, la plus forte erreur apparaît sur "
            f"{worst_lstm['ticker']} le {worst_lstm['date']} ; pour le GRU, "
            f"elle apparaît sur {worst_gru['ticker']} le {worst_gru['date']}. "
            "Ces dates correspondent à des hausses brutales de High - Low, donc "
            "à des épisodes de volatilité que les retards passés expliquent mal."
        ),
        (
            "L'un des deux modèles capture-t-il mieux les pics de volatilité ? "
            "Les deux modèles capturent la dynamique générale mais sous-estiment "
            f"les pics. Dans cette expérience, {best['model']} présente un léger "
            f"avantage global (RMSE {best['rmse']:.4f}, R² {best['r2']:.4f}) "
            f"et {fastest['model']} est le plus rapide à entraîner."
        ),
        (
            "Conclusion sur la capacité des modèles : les modèles LSTM et GRU "
            "sont utiles pour anticiper le régime moyen de volatilité journalière, "
            "mais leur capacité à prévoir les chocs extrêmes reste limitée sans "
            "variables explicatives supplémentaires comme des indicateurs de "
            "marché, des volumes ou des événements macroéconomiques."
        ),
    ]
    add_bullets(doc, answers)


def add_stats_table(doc: Document, summary: pd.DataFrame) -> None:
    """
    Add an excerpt of ticker-level statistics

    :param doc: Word document
    :param summary: ticker summary
    """
    sample = summary.head(5)
    table = doc.add_table(rows=1, cols=6)
    set_table_width(table, [0.8, 1.2, 1.2, 1.1, 1.1, 1.1])
    set_table_borders(table)
    headers = ["Ticker", "Prix moyen", "Rend. moyen", "Variance", "Volatilité", "Bêta"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, LIGHT_GRAY)
        set_run_font(cell.paragraphs[0].runs[0], bold=True)

    for _, row_data in sample.iterrows():
        row = table.add_row()
        values = [
            row_data["ticker"],
            f"{row_data['mean_price']:.4f}",
            f"{row_data['mean_return']:.6f}",
            f"{row_data['variance']:.6f}",
            f"{row_data['volatility']:.6f}",
            f"{row_data['beta']:.4f}",
        ]
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def add_picture_with_caption(
    doc: Document,
    image_path: Path,
    caption: str,
    width: float = 5.9,
) -> None:
    """
    Add a figure and caption if the image exists

    :param doc: Word document
    :param image_path: figure path
    :param caption: caption text
    :param width: image width in inches
    """
    if not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(10)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=MUTED, italic=True)


def build_report() -> None:
    """
    Build the complete Word report from the CSV outputs in the outputs directory
    """
    metrics = pd.read_csv(OUTPUT_DIR / "model_metrics.csv")
    summary = pd.read_csv(OUTPUT_DIR / "ticker_summary.csv")
    predictions = pd.read_csv(OUTPUT_DIR / "predictions_test.csv")

    best = metrics.sort_values("rmse").iloc[0]
    fastest = metrics.sort_values("training_time_seconds").iloc[0]

    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    doc.add_heading("1. Introduction", level=1)
    add_paragraph(
        doc,
        "Ce projet se place dans le cadre d'une analyse de marché appliquée aux "
        "actifs du CAC 40. L'objectif est de construire un jeu de données clair, "
        "de calculer des indicateurs financiers par actif, puis de comparer deux "
        "modèles séquentiels, LSTM et GRU, pour prédire la volatilité intraday "
        "mesurée par l'écart High - Low.",
    )
    add_paragraph(
        doc,
        "Le code a été organisé en programmation orientée objet afin de séparer "
        "la lecture des données, l'analyse statistique, la construction des "
        "séquences temporelles, l'entraînement des modèles et la visualisation "
        "des résultats."
    )

    doc.add_heading("2. Données et préparation", level=1)
    add_paragraph(
        doc,
        "Le fichier cac40.csv contient les prix d'ouverture, plus haut, plus bas, "
        "clôture, volume échangé et informations de date pour 40 actifs. La "
        "fonction read_data lit le fichier avec pandas et retourne un objet "
        "pd.DataFrame. La contrainte d'un seul appel à pd.read_csv est respectée."
    )
    add_bullets(
        doc,
        [
            "Les colonnes jour, mois et annee sont renommées en DAY, MONTH et YEAR.",
            "Une colonne DATE est créée pour obtenir une vraie série temporelle.",
            "Les observations sont triées par TICKER puis DATE.",
            "Les données sont regroupées par actif dans un dictionnaire par ticker.",
        ],
    )

    doc.add_heading("3. Organisation orientée objet", level=1)
    add_paragraph(
        doc,
        "La structure du code repose sur plusieurs classes. MarketAnalyzer calcule "
        "les rendements, les statistiques et les bêtas. SequenceDatasetBuilder "
        "transforme les séries High - Low en séquences retardées. "
        "RecurrentVolatilityModel encapsule la construction et l'entraînement "
        "des modèles LSTM et GRU. ResultVisualizer produit les graphiques de "
        "comparaison."
    )
    add_bullets(
        doc,
        [
            "Chaque classe et chaque fonction importante possède une docstring.",
            "Les entrées et sorties sont typées avec les annotations Python.",
            "Les responsabilités sont séparées afin de rendre le code lisible et maintenable.",
        ],
    )

    doc.add_heading("4. Statistiques et bêta", level=1)
    add_paragraph(
        doc,
        "Pour chaque actif, le programme calcule le prix moyen de clôture, le "
        "rendement moyen, la variance des rendements et la volatilité des "
        "rendements. Ces valeurs sont stockées dans les dictionnaires "
        "mean_price_by_ticker, mean_return_by_ticker, variance_by_ticker et "
        "volatility_by_ticker."
    )
    add_paragraph(
        doc,
        "Le rendement de marché rm,t est construit comme la moyenne simple des "
        "rendements de tous les actifs à chaque date. Le bêta de chaque actif est "
        "ensuite estimé par la régression ri,t = alpha_i + beta_i rm,t. Le taux "
        "sans risque est fixé à 0,01 et l'alpha de Jensen est calculé à partir de "
        "l'ordonnée à l'origine."
    )
    add_paragraph(doc, "Extrait des résultats par actif :")
    add_stats_table(doc, summary)

    doc.add_heading("5. Modélisation LSTM et GRU", level=1)
    add_paragraph(
        doc,
        "La variable cible de la partie séquentielle est la volatilité intraday "
        "HLi,t = Highi,t - Lowi,t. Pour chaque actif, le programme construit une "
        "séquence de p retards : Xi,t = [HLi,t-1, HLi,t-2, ..., HLi,t-p]. Les "
        "données sont ensuite reformattées en tenseurs de forme (p, 1), ce qui "
        "correspond à l'entrée attendue par les réseaux récurrents."
    )
    add_bullets(
        doc,
        [
            "Train : 70 % des séquences, en respectant l'ordre temporel.",
            "Validation : 15 % des séquences.",
            "Test : 15 % des séquences.",
            "Les modèles sont évalués avec le RMSE et le coefficient R².",
        ],
    )
    add_metrics_table(doc, metrics)

    doc.add_heading("6. Visualisation et interprétation", level=1)
    add_paragraph(
        doc,
        "Les modèles LSTM et GRU produisent des prédictions sur l'ensemble de "
        "test. Les graphiques comparent les valeurs réelles de High - Low aux "
        "valeurs prédites. Les deux modèles suivent correctement le niveau moyen "
        "de volatilité, mais ils lissent les pics extrêmes. Ce comportement est "
        "classique dans une prévision utilisant uniquement des retards passés : "
        "les chocs de marché soudains sont difficiles à anticiper avant leur "
        "apparition dans la série."
    )
    add_picture_with_caption(
        doc,
        OUTPUT_DIR / "predictions_AC.png",
        "Figure 1 - Comparaison des valeurs réelles et prédites pour le titre AC.",
    )
    add_picture_with_caption(
        doc,
        OUTPUT_DIR / "scatter_predictions.png",
        "Figure 2 - Nuage de points High - Low réel contre High - Low prédit.",
        width=4.8,
    )
    add_paragraph(
        doc,
        f"Dans l'exécution finale, le meilleur modèle selon le RMSE est "
        f"{best['model']}, avec un RMSE de {best['rmse']:.4f} et un R² de "
        f"{best['r2']:.4f}. Le modèle le plus rapide à entraîner est "
        f"{fastest['model']}, avec un temps d'entraînement de "
        f"{fastest['training_time_seconds']:.2f} secondes."
    )

    add_explicit_answers(doc, metrics, predictions)

    doc.add_heading("8. Difficultés rencontrées", level=1)
    add_paragraph(
        doc,
        "La première difficulté concerne le format du fichier csv : les décimales "
        "sont écrites avec une virgule et les colonnes de date sont séparées en "
        "jour, mois et année. Il a donc fallu nettoyer les types numériques et "
        "reconstruire une colonne DATE exploitable."
    )
    add_paragraph(
        doc,
        "La seconde difficulté est liée à la modélisation des séries financières. "
        "La volatilité intraday peut présenter des pics très brusques. Comme les "
        "modèles utilisent seulement les valeurs passées de High - Low, ils "
        "capturent bien la tendance générale mais ont tendance à sous-estimer les "
        "chocs exceptionnels."
    )
    add_paragraph(
        doc,
        "Enfin, le coût de calcul des modèles récurrents est non négligeable. Le "
        "GRU obtient ici une performance très proche de celle du LSTM avec un "
        "temps d'entraînement inférieur, ce qui en fait un modèle intéressant "
        "pour ce problème."
    )

    doc.add_heading("9. Conclusion", level=1)
    add_paragraph(
        doc,
        "Le projet montre qu'une approche orientée objet permet d'organiser "
        "clairement une analyse financière complète : lecture des données, "
        "statistiques, estimation des bêtas, construction de séquences, "
        "entraînement des modèles et visualisation. Les résultats indiquent que "
        "les réseaux LSTM et GRU peuvent anticiper le niveau moyen de volatilité "
        f"journalière, mais restent limités face aux pics soudains. Dans cette "
        f"expérience, {best['model']} obtient le meilleur RMSE ({best['rmse']:.4f}) "
        f"et {fastest['model']} est le plus rapide à entraîner "
        f"({fastest['training_time_seconds']:.1f} s)."
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_report()
